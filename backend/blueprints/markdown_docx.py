import traceback
from io import BytesIO

from flask import Blueprint, request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from lxml import html as lxml_html

from utils.helpers import error, send_file_and_cleanup

try:
    import markdown2
except ImportError:
    markdown2 = None

markdown_docx_bp = Blueprint("markdown_docx", __name__)


def _add_inline(paragraph, element):
    if element.text:
        paragraph.add_run(element.text)

    for child in element:
        tag = child.tag if isinstance(child.tag, str) else ""

        if tag in ("strong", "b"):
            run = paragraph.add_run(child.text or "")
            run.bold = True
        elif tag in ("em", "i"):
            run = paragraph.add_run(child.text or "")
            run.italic = True
        elif tag == "code":
            run = paragraph.add_run(child.text or "")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif tag == "a":
            run = paragraph.add_run(child.text or "")
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
        elif tag == "br":
            paragraph.add_run("\n")
        else:
            _add_inline(paragraph, child)

        if child.tail:
            paragraph.add_run(child.tail)


def _process_table(doc, table_elem):
    rows = table_elem.findall(".//tr")
    if not rows:
        return

    cols = 0
    for row in rows:
        cells = row.findall("th") or row.findall("td")
        cols = max(cols, len(cells))

    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        cells = row.findall("th") or row.findall("td")
        is_header = bool(row.findall("th"))
        for j, cell in enumerate(cells):
            if j >= cols:
                break
            table.cell(i, j).text = cell.text_content().strip()
            if is_header:
                for p in table.cell(i, j).paragraphs:
                    for run in p.runs:
                        run.bold = True


def _convert(md_text):
    html_content = markdown2.markdown(
        md_text, extras=["fenced-code-blocks", "tables"]
    )
    root = lxml_html.fromstring(f"<div>{html_content}</div>")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for el in root:
        tag = el.tag if isinstance(el.tag, str) else ""

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            doc.add_heading(el.text_content().strip(), level=level)

        elif tag == "p":
            p = doc.add_paragraph()
            _add_inline(p, el)

        elif tag in ("ul", "ol"):
            style_name = "List Bullet" if tag == "ul" else "List Number"
            for li in el:
                if li.tag == "li":
                    p = doc.add_paragraph(style=style_name)
                    _add_inline(p, li)

        elif tag == "pre":
            code = el.find("code")
            text = code.text_content() if code is not None else el.text_content()
            p = doc.add_paragraph()
            run = p.add_run(text.rstrip("\n"))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.3)

        elif tag == "blockquote":
            for child in el:
                if child.tag == "p":
                    p = doc.add_paragraph()
                    _add_inline(p, child)
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.italic = True

        elif tag == "hr":
            p = doc.add_paragraph()
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(200, 200, 200)

        elif tag == "table":
            _process_table(doc, el)

        else:
            text = el.text_content().strip()
            if text:
                doc.add_paragraph(text)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@markdown_docx_bp.route("/convertMdToDocx", methods=["POST"])
def convert_md_to_docx():
    if markdown2 is None:
        return error("markdown2 dependency is not installed on the server", 500)

    try:
        if "file" not in request.files:
            return error("No file provided")

        md_file = request.files["file"]
        if md_file.filename == "":
            return error("No file selected")
        if not md_file.filename.lower().endswith(".md"):
            return error("Invalid file format. Please upload a Markdown (.md) file.")

        md_content = md_file.read().decode("utf-8")
        docx_buffer = _convert(md_content)

        base = md_file.filename.rsplit(".", 1)[0]
        download_name = base + ".docx"

        return send_file_and_cleanup(
            docx_buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=download_name,
        )

    except UnicodeDecodeError:
        return error(
            "Could not decode the file as UTF-8. Please ensure it is a valid Markdown file."
        )
    except Exception as e:
        traceback.print_exc()
        return error("An error occurred during conversion: " + str(e), 500)
